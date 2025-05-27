"""
超星学习通作业爬取工具
支持登录、课程选择、作业爬取和导出功能
"""

import base64
import json
import logging
import re
import time
from argparse import ArgumentParser
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
from urllib import parse
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from bs4.element import Tag
from Crypto.Cipher import AES
from Crypto.Util.Padding import pad, unpad


# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('fanya_crawler.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


@dataclass
class Question:
    """题目数据类"""
    answer_type: int
    question_title: str
    correct_answer: Union[str, List[str]]
    question_answers: Optional[List[str]] = None


@dataclass
class Assignment:
    """作业数据类"""
    work_id: str
    assignment_name: str
    assignment_status: str
    assignment_url: str
    course_id: str
    questions: List[Question] = field(default_factory=list)


@dataclass
class Course:
    """课程数据类"""
    course_id: str
    class_id: str
    cpi: str
    course_name: str
    course_url: str


class FanyaCrawlerError(Exception):
    """自定义异常类"""
    pass


class AESCrypto:
    """AES加密解密工具类"""

    def __init__(self, key: str = "u2oh6Vu^HWe4_AES"):
        self.key = key.encode('utf-8')

    def encrypt(self, message: str) -> str:
        """AES加密"""
        try:
            cipher = AES.new(self.key, AES.MODE_CBC, self.key)
            padded_message = pad(message.encode('utf-8'),
                                 AES.block_size, style='pkcs7')
            encrypted_bytes = cipher.encrypt(padded_message)
            return base64.b64encode(encrypted_bytes).decode('utf-8')
        except Exception as e:
            logger.error(f"加密失败: {e}")
            raise FanyaCrawlerError(f"加密失败: {e}")

    def decrypt(self, encrypted_str: str) -> Tuple[bool, str]:
        """AES解密"""
        try:
            encrypted_bytes = base64.b64decode(encrypted_str)
            cipher = AES.new(self.key, AES.MODE_CBC, self.key)
            decrypted_bytes = unpad(cipher.decrypt(
                encrypted_bytes), AES.block_size)
            return True, decrypted_bytes.decode('utf-8')
        except Exception as e:
            logger.error(f"解密失败: {e}")
            return False, str(e)


class FanyaCrawler:
    """超星学习通爬虫主类"""

    # 答题类型映射
    ANSWER_TYPES = {
        "单选题": 1,
        "多选题": 2,
        "填空题": 3,
        "判断题": 4,
        "思维导图": 5,
        "其它": 255
    }

    # API endpoints
    API_ENDPOINTS = {
        'login': 'http://passport2.chaoxing.com/fanyalogin',
        'course_list': 'https://mooc2-ans.chaoxing.com/mooc2-ans/visit/courselistdata',
        'course_middle': 'https://mooc1.chaoxing.com/visit/stucoursemiddle',
        'work_list': 'https://mooc1.chaoxing.com/mooc2/work/list',
        'work_view': 'https://mooc1.chaoxing.com/mooc-ans/mooc2/work/view'
    }

    def __init__(self):
        self.session = requests.Session()
        self.crypto = AESCrypto()
        self._setup_session()

    def _setup_session(self):
        """配置会话"""
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
            '(KHTML, like Gecko) Chrome/132.0.0.0 Safari/537.36 Edg/132.0.0.0'
        })

    def login(self, phone: str, password: str) -> bool:
        """登录功能"""
        try:
            encrypted_phone = self.crypto.encrypt(phone)
            encrypted_password = self.crypto.encrypt(password)

            headers = {
                "Accept": "application/json, text/javascript, */*; q=0.01",
                "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
                "Origin": "https://passport2.chaoxing.com",
                "Referer": "https://passport2.chaoxing.com/login",
                "X-Requested-With": "XMLHttpRequest"
            }

            payload = {
                'uname': encrypted_phone,
                'password': encrypted_password,
                't': "true",
                'validate': None,
                'doubleFactorLogin': "0",
                'independentId': "0"
            }

            response = self.session.post(
                self.API_ENDPOINTS['login'],
                data=payload,
                headers=headers,
                timeout=30
            )
            response.raise_for_status()

            result = response.json()
            if result.get('status'):
                logger.info("登录成功")
                return True
            else:
                logger.error(f"登录失败: {result.get('msg', '未知错误')}")
                return False

        except requests.RequestException as e:
            logger.error(f"登录请求失败: {e}")
            return False
        except Exception as e:
            logger.error(f"登录异常: {e}")
            return False

    def get_courses(self) -> List[Course]:
        """获取课程列表"""
        try:
            params = {
                "courseType": "1",
                "courseFolderId": "0",
                "superstarClass": "0"
            }

            response = self.session.get(
                self.API_ENDPOINTS['course_list'],
                params=params,
                timeout=30
            )
            response.raise_for_status()

            soup = BeautifulSoup(response.text, "lxml")
            course_divs = soup.find_all("div", attrs={"class": "course-info"})

            courses = []
            for course_div in course_divs:
                try:
                    link = course_div.find_next("a", attrs={"class": "color1"})
                    if not link:
                        continue

                    course_url = link.get("href", "")
                    name_span = link.find_next(
                        "span", attrs={"class": "course-name overHidden2"})
                    if not name_span:
                        continue

                    course_name = name_span.get_text(strip=True)

                    # 解析URL参数
                    parsed_url = urlparse(course_url)
                    url_data = parse.parse_qs(parsed_url.query)

                    course = Course(
                        course_id=url_data.get("courseid", [""])[0],
                        class_id=url_data.get("clazzid", [""])[0],
                        cpi=url_data.get("cpi", [""])[0],
                        course_name=course_name,
                        course_url=course_url
                    )

                    if course.course_id and course.class_id:
                        courses.append(course)

                except Exception as e:
                    logger.warning(f"解析课程信息失败: {e}")
                    continue

            logger.info(f"获取到 {len(courses)} 门课程")
            return courses

        except Exception as e:
            logger.error(f"获取课程列表失败: {e}")
            raise FanyaCrawlerError(f"获取课程列表失败: {e}")

    def get_assignments(self, course: Course) -> List[Assignment]:
        """获取课程作业列表"""
        try:
            # 获取课程中间页面
            middle_params = {
                "courseid": course.course_id,
                "clazzid": course.class_id,
                "cpi": course.cpi,
                "ismooc2": 1,
                "v": time.time(),
                "start": 0,
                "size": 500,
                "catalogId": 0,
                "superstarClass": 0,
            }

            middle_response = self.session.get(
                self.API_ENDPOINTS['course_middle'],
                params=middle_params,
                timeout=30
            )
            middle_response.raise_for_status()

            # 提取作业编码参数
            soup = BeautifulSoup(middle_response.text, "lxml")
            work_enc_input = soup.find("input", id="workEnc")
            if not work_enc_input:
                logger.error("未找到工作加密参数")
                return []

            work_enc = work_enc_input.get("value", "")

            # 获取作业列表
            work_params = {
                "courseId": course.course_id,
                "classId": course.class_id,
                "cpi": course.cpi,
                "ut": "s",
                "enc": work_enc,
            }

            assignments = []
            page_num = 1
            total_page = 1

            while True:
                work_params["pageNum"] = page_num

                response = self.session.get(
                    self.API_ENDPOINTS['work_list'],
                    params=work_params,
                    timeout=30
                )
                response.raise_for_status()

                soup = BeautifulSoup(response.text, "lxml")

                if page_num == 1:
                    pagenum_pattern = r'pageNum\s*:\s*(\d+)'
                    scripts_in_soup = soup.find("body").find("script")
                    print(scripts_in_soup)
                    jscode_in_soup = scripts_in_soup.string if scripts_in_soup else ""
                    total_pagenum = re.search(pagenum_pattern, jscode_in_soup)

                    if (total_pagenum):
                        total_page = total_pagenum.group(1)

                li_tags = soup.find_all("li")

                if not li_tags:
                    break

                page_assignments = []
                for li in li_tags:
                    try:
                        data_url = li.get("data", "")
                        if not data_url:
                            continue

                        parsed_url = urlparse(data_url)
                        url_data = parse.parse_qs(parsed_url.query)

                        work_id = url_data.get("workId", [""])[0]
                        if not work_id:
                            continue

                        # 获取作业名称和状态
                        name_p = li.find("p")
                        assignment_name = name_p.get_text(
                            strip=True) if name_p else "未知作业"

                        status_p = name_p.find_next("p") if name_p else None
                        assignment_status = status_p.get_text(
                            strip=True) if status_p else "未知状态"

                        assignment = Assignment(
                            work_id=work_id,
                            assignment_name=assignment_name,
                            assignment_status=assignment_status,
                            assignment_url=data_url,
                            course_id=course.course_id
                        )

                        page_assignments.append(assignment)

                    except Exception as e:
                        logger.warning(f"解析作业信息失败: {e}")
                        continue

                assignments.extend(page_assignments)

                # 检查是否还有下一页
                if (total_page == page_num):
                    break

                if (total_page != 1):
                    page_num += 1

            logger.info(f"获取到 {len(assignments)} 个作业")
            return assignments

        except Exception as e:
            logger.error(f"获取作业列表失败: {e}")
            raise FanyaCrawlerError(f"获取作业列表失败: {e}")

    def _normalize_title(self, title: str) -> str:
        """标准化题目标题"""
        return (title.strip()
                .replace("（", "(")
                .replace("）", ")")
                .replace("\u200c", "")
                .replace("\u200e", "")
                .replace("\u200d", "")
                .replace("\u200f", "")
                .replace("\xa0", ""))

    def _normalize_answers(self, answers_text: str) -> List[str]:
        """标准化答案选项"""
        answer_list = answers_text.split("\n")
        return [answer.strip() + "\n" for answer in answer_list if answer.strip()]

    def _parse_questions(self, question_block: Tag) -> List[Question]:
        """解析题目块"""
        try:
            title_tag = question_block.find("h2", attrs={"class": "type_tit"})
            if not title_tag:
                return []

            block_title = title_tag.get_text(strip=True)

            # 确定题目类型
            answer_type = 0
            for type_name, type_id in self.ANSWER_TYPES.items():
                if type_name in block_title:
                    answer_type = type_id
                    break

            if answer_type == 0:
                logger.warning(f"未知题目类型: {block_title}")
                return []

            logger.info(f"解析题目类型: {block_title}")

            # 获取所有题目详情
            question_details = question_block.find_all(
                "div", attrs={"aria-label": "题目详情"})
            questions = []

            for detail in question_details:
                try:
                    question = self._parse_single_question(detail, answer_type)
                    if question:
                        questions.append(question)
                except Exception as e:
                    logger.warning(f"解析单个题目失败: {e}")
                    continue

            return questions

        except Exception as e:
            logger.error(f"解析题目块失败: {e}")
            return []

    def _parse_single_question(self, detail: Tag, answer_type: int) -> Optional[Question]:
        """解析单个题目"""
        try:
            # 获取题目标题
            title_tag = detail.find(
                "h3", attrs={"class": "mark_name colorDeep"})
            if not title_tag:
                return None

            question_title = self._normalize_title(title_tag.get_text())

            if answer_type in [self.ANSWER_TYPES["单选题"], self.ANSWER_TYPES["多选题"]]:
                # 选择题
                answers_tag = detail.find(
                    "ul", attrs={"class": "mark_letter colorDeep qtDetail"})
                question_answers = self._normalize_answers(
                    answers_tag.get_text()) if answers_tag else []

                answer_div = detail.find("div", attrs={"class": "mark_answer"})
                correct_tag = answer_div.find(
                    "span", attrs={"class": "rightAnswerContent workTextWrap"}) if answer_div else None
                correct_answer = correct_tag.get_text(
                    strip=True) if correct_tag else ""

                return Question(
                    answer_type=answer_type,
                    question_title=question_title,
                    question_answers=question_answers,
                    correct_answer=correct_answer
                )

            elif answer_type == self.ANSWER_TYPES["填空题"]:
                # 填空题
                fill_tag = detail.find(
                    "dl", attrs={"class": "mark_fill colorGreen"})
                if fill_tag:
                    dd_tags = fill_tag.find_all("dd")
                    correct_answers = [dd.get_text(
                        strip=True) for dd in dd_tags]
                else:
                    correct_answers = []

                return Question(
                    answer_type=answer_type,
                    question_title=question_title,
                    correct_answer=correct_answers
                )

            elif answer_type == self.ANSWER_TYPES["判断题"]:
                # 判断题
                answer_div = detail.find("div", attrs={"class": "mark_answer"})
                correct_tag = answer_div.find(
                    "span", attrs={"class": "rightAnswerContent"}) if answer_div else None
                correct_answer = correct_tag.get_text(
                    strip=True) if correct_tag else ""

                return Question(
                    answer_type=answer_type,
                    question_title=question_title,
                    correct_answer=correct_answer
                )

        except Exception as e:
            logger.warning(f"解析单个题目详情失败: {e}")

        return None

    def get_assignment_questions(self, assignment: Assignment) -> List[Question]:
        """获取作业题目"""
        try:
            parsed_url = urlparse(assignment.assignment_url)
            url_data = parse.parse_qs(parsed_url.query)

            params = {
                "courseId": url_data.get("courseId", [""])[0],
                "classId": url_data.get("classId", [""])[0],
                "cpi": url_data.get("cpi", [""])[0],
                "workId": url_data.get("workId", [""])[0],
                "answerId": url_data.get("answerId", [""])[0],
                "enc": url_data.get("enc", [""])[0]
            }

            response = self.session.get(
                self.API_ENDPOINTS['work_view'],
                params=params,
                timeout=30
            )
            response.raise_for_status()

            soup = BeautifulSoup(response.text, "lxml")
            question_blocks = soup.find_all(
                "div", attrs={"class": "mark_item"})

            all_questions = []
            for block in question_blocks:
                questions = self._parse_questions(block)
                all_questions.extend(questions)

            logger.info(
                f"作业 {assignment.assignment_name} 获取到 {len(all_questions)} 道题目")
            return all_questions

        except Exception as e:
            logger.error(f"获取作业题目失败: {e}")
            return []


class DocumentExporter:
    """文档导出器"""

    def __init__(self, course_name: str):
        self.course_name = course_name
        self.output_dir = Path("output")
        self.output_dir.mkdir(exist_ok=True)

    def export_markdown(self, assignments: List[Assignment], with_answers: bool = True):
        """导出Markdown格式"""
        suffix = "带答案" if with_answers else "不带答案"
        filename = self.output_dir / f"{self.course_name}_习题_{suffix}.md"

        try:
            with open(filename, "w", encoding="utf-8") as f:
                f.write(f"# {self.course_name} 习题集\n\n")

                for assignment in assignments:
                    if not assignment.questions:
                        continue

                    f.write(f"## {assignment.assignment_name}\n\n")

                    for question in assignment.questions:
                        f.write(f"### {question.question_title}\n\n")

                        # 选择题显示选项
                        if question.question_answers:
                            for answer in question.question_answers:
                                f.write(answer)
                            f.write("\n")

                        # 显示答案或留空
                        if with_answers:
                            if isinstance(question.correct_answer, list):
                                f.write(
                                    "正确答案: " + ", ".join(question.correct_answer) + "\n\n")
                            else:
                                f.write(f"正确答案: {question.correct_answer}\n\n")
                        else:
                            f.write("答案: ____________________\n\n")

            logger.info(f"Markdown导出完成: {filename}")

        except Exception as e:
            logger.error(f"Markdown导出失败: {e}")

    def export_word(self, assignments: List[Assignment], with_answers: bool = True):
        """导出Word格式"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.oxml.ns import qn
        except ImportError:
            logger.error("python-docx库未安装，无法导出Word文档")
            return

        suffix = "带答案" if with_answers else "不带答案"
        filename = self.output_dir / f"{self.course_name}_习题_{suffix}.docx"

        try:
            doc = Document()

            # 设置默认字体
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 添加标题
            doc.add_heading(f"{self.course_name} 习题集", level=0)

            for assignment in assignments:
                if not assignment.questions:
                    continue

                # 作业标题
                doc.add_heading(assignment.assignment_name, level=1)

                for i, question in enumerate(assignment.questions, 1):
                    # 题目标题
                    para = doc.add_paragraph()
                    run = para.add_run(f"{i}. {question.question_title}")
                    run.bold = True

                    # 选择题选项
                    if question.question_answers:
                        for answer in question.question_answers:
                            p = doc.add_paragraph(
                                answer.strip(), style="List Bullet")
                            p.paragraph_format.left_indent = Inches(0.5)

                    # 答案部分
                    if with_answers:
                        ans_para = doc.add_paragraph()
                        if isinstance(question.correct_answer, list):
                            ans_text = "正确答案: " + \
                                ", ".join(question.correct_answer)
                        else:
                            ans_text = f"正确答案: {question.correct_answer}"
                        run = ans_para.add_run(ans_text)
                        run.bold = True
                    else:
                        doc.add_paragraph("答案: ____________________")

                    # 添加空行
                    doc.add_paragraph()

            doc.save(filename)
            logger.info(f"Word导出完成: {filename}")

        except Exception as e:
            logger.error(f"Word导出失败: {e}")

    def export_json(self, assignments: List[Assignment]):
        """导出JSON格式（用于备份和数据交换）"""
        filename = self.output_dir / f"{self.course_name}_习题_数据.json"

        try:
            data = {
                "course_name": self.course_name,
                "export_time": time.strftime("%Y-%m-%d %H:%M:%S"),
                "assignments": []
            }

            for assignment in assignments:
                assignment_data = {
                    "work_id": assignment.work_id,
                    "name": assignment.assignment_name,
                    "status": assignment.assignment_status,
                    "questions": []
                }

                for question in assignment.questions:
                    question_data = {
                        "type": question.answer_type,
                        "title": question.question_title,
                        "correct_answer": question.correct_answer,
                        "options": question.question_answers
                    }
                    assignment_data["questions"].append(question_data)

                data["assignments"].append(assignment_data)

            with open(filename, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            logger.info(f"JSON导出完成: {filename}")

        except Exception as e:
            logger.error(f"JSON导出失败: {e}")


def main():
    """主函数"""
    parser = ArgumentParser(
        prog="fanya_crawler",
        description="超星学习通作业爬取工具"
    )
    parser.add_argument("phone", help="手机号")
    parser.add_argument("password", help="密码")
    parser.add_argument("--format", choices=["markdown", "word", "json", "all"],
                        default="all", help="导出格式")
    parser.add_argument("--no-answers", action="store_true", help="不包含答案")

    args = parser.parse_args()

    try:
        # 初始化爬虫
        crawler = FanyaCrawler()

        # 登录
        logger.info("开始登录...")
        if not crawler.login(args.phone, args.password):
            logger.error("登录失败，程序退出")
            return

        # 获取课程列表
        logger.info("获取课程列表...")
        courses = crawler.get_courses()
        if not courses:
            logger.error("未找到课程，程序退出")
            return

        # 显示课程列表供用户选择
        print("\n请选择要爬取的课程:")
        for i, course in enumerate(courses, 1):
            print(f"{i}. {course.course_name}")

        while True:
            try:
                choice = int(input(f"\n请输入课程编号 (1-{len(courses)}): "))
                if 1 <= choice <= len(courses):
                    selected_course = courses[choice - 1]
                    break
                else:
                    print(f"请输入 1 到 {len(courses)} 之间的数字")
            except ValueError:
                print("请输入有效的数字")

        logger.info(f"选择课程: {selected_course.course_name}")

        # 获取作业列表
        logger.info("获取作业列表...")
        assignments = crawler.get_assignments(selected_course)
        if not assignments:
            logger.warning("该课程暂无作业")
            return

        # 获取每个作业的题目
        logger.info("开始爬取作业题目...")
        for assignment in assignments:
            logger.info(f"正在处理作业: {assignment.assignment_name}")
            questions = crawler.get_assignment_questions(assignment)
            assignment.questions = questions
            time.sleep(1)  # 避免请求过于频繁

        # 导出文档
        exporter = DocumentExporter(selected_course.course_name)

        if args.format in ["markdown", "all"]:
            exporter.export_markdown(
                assignments, with_answers=not args.no_answers)
            if not args.no_answers:
                exporter.export_markdown(assignments, with_answers=False)

        if args.format in ["word", "all"]:
            exporter.export_word(assignments, with_answers=not args.no_answers)
            if not args.no_answers:
                exporter.export_word(assignments, with_answers=False)

        if args.format in ["json", "all"]:
            exporter.export_json(assignments)

        logger.info("所有任务完成！")

    except KeyboardInterrupt:
        logger.info("用户中断程序")
    except Exception as e:
        logger.error(f"程序执行出错: {e}")
        raise


if __name__ == "__main__":
    main()
